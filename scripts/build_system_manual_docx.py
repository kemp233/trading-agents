from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "system-manual.md"
OUTPUT = ROOT / "output" / "doc" / "Trading_Agents_CTP_System_Manual_Submission_FINAL.docx"


def set_run_font(run, size: int, bold: bool = False, east_asia: str = "SimSun", latin: str = "Times New Roman") -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    section.start_type = WD_SECTION.NEW_PAGE


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    for style_name, east_asia in [("Heading 1", "SimHei"), ("Heading 2", "SimHei"), ("Heading 3", "SimHei")]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_spacing(paragraph, line_spacing: float = 1.5, before: int = 0, after: int = 6, first_line_cm: float = 0.0) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if first_line_cm:
        fmt.first_line_indent = Cm(first_line_cm)


def add_code_block(doc: Document, lines: list[str]) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, line_spacing=1.2, after=6)
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, 10, east_asia="Consolas", latin="Consolas")
        if idx < len(lines) - 1:
            run.add_break()

    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F4F4F4")
    p_pr.append(shading)


def add_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for left, right in rows:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
        for cell in row:
            for p in cell.paragraphs:
                set_paragraph_spacing(p, line_spacing=1.25, after=0)
                for run in p.runs:
                    set_run_font(run, 10.5)


def render_markdown(doc: Document, markdown_text: str) -> None:
    lines = markdown_text.splitlines()
    in_code = False
    code_lines: list[str] = []
    info_rows: list[tuple[str, str]] = []
    in_info_block = False

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.strip() == "```text":
            in_code = True
            code_lines = []
            continue

        if line.strip() == "```" and in_code:
            add_code_block(doc, code_lines)
            in_code = False
            code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line[2:].lstrip("\ufeff"))
            set_run_font(run, 18, bold=True, east_asia="SimHei")
            set_paragraph_spacing(paragraph, line_spacing=1.2, after=12)
            continue

        if line.startswith("## "):
            if in_info_block and info_rows:
                add_table(doc, info_rows)
                info_rows = []
                in_info_block = False
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line[3:])
            set_run_font(run, 14, bold=True, east_asia="SimHei")
            set_paragraph_spacing(paragraph, line_spacing=1.3, before=6, after=6)
            continue

        if line.startswith("### "):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line[4:])
            set_run_font(run, 12.5, bold=True, east_asia="SimHei")
            set_paragraph_spacing(paragraph, line_spacing=1.3, before=4, after=4)
            continue

        if not line.strip():
            if in_info_block and info_rows:
                add_table(doc, info_rows)
                info_rows = []
                in_info_block = False
            continue

        if line.startswith("- "):
            content = line[2:]
            if "：" in content and not content.startswith("截至 "):
                left, right = content.split("：", 1)
                if not in_info_block:
                    in_info_block = True
                info_rows.append((left, right.replace("`", "")))
            else:
                if in_info_block and info_rows:
                    add_table(doc, info_rows)
                    info_rows = []
                    in_info_block = False
                paragraph = doc.add_paragraph(style="List Bullet")
                run = paragraph.add_run(content.replace("`", ""))
                set_run_font(run, 11)
                set_paragraph_spacing(paragraph, line_spacing=1.35, after=3)
            continue

        if in_info_block and info_rows:
            add_table(doc, info_rows)
            info_rows = []
            in_info_block = False

        paragraph = doc.add_paragraph()
        run = paragraph.add_run(line.replace("`", ""))
        set_run_font(run, 11)
        set_paragraph_spacing(paragraph, line_spacing=1.5, after=6, first_line_cm=0.74)

    if in_info_block and info_rows:
        add_table(doc, info_rows)


def build() -> Path:
    markdown_text = SOURCE.read_text(encoding="utf-8-sig")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    render_markdown(doc, markdown_text)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build()
    print(output)
